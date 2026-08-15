"""
InsightIQ — AI Business Decision Engine
Multi-provider AI · Groq · OpenAI · Gemini
All fixes applied: Gemini v2, bar chart, tab syntax, professional exports
"""

import os, io, re, json, warnings, datetime
import pandas as pd
import streamlit as st
from dotenv import load_dotenv

warnings.filterwarnings("ignore")
load_dotenv()

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="InsightIQ · AI Business Analyst",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --bg:#070B14; --surface:#0F1624; --surface2:#151D2E; --surface3:#1A2338;
  --border:#1E2D45; --border2:#243550;
  --accent:#3B82F6; --accent2:#6366F1; --accent3:#8B5CF6;
  --success:#10B981; --warning:#F59E0B; --danger:#EF4444;
  --text:#F1F5F9; --text2:#CBD5E1; --muted:#64748B; --muted2:#475569;
  --r:16px; --r-sm:10px;
  --grad:linear-gradient(135deg,#3B82F6,#6366F1,#8B5CF6);
  --glow:0 0 40px rgba(59,130,246,.15);
  --shadow:0 8px 32px rgba(0,0,0,.4);
}
*,*::before,*::after{box-sizing:border-box;}
html,body,[class*="css"],.stApp{
  font-family:'Inter',-apple-system,sans-serif!important;
  background:var(--bg)!important;color:var(--text)!important;
  -webkit-font-smoothing:antialiased;
}
#MainMenu,footer,header,[data-testid="stToolbar"],
[data-testid="stDecoration"],[data-testid="stStatusWidget"]{display:none!important;}
::-webkit-scrollbar{width:3px;height:3px;}
::-webkit-scrollbar-thumb{background:var(--border2);border-radius:3px;}

@keyframes fadeUp{from{opacity:0;transform:translateY(16px)}to{opacity:1;transform:translateY(0)}}
@keyframes fadeIn{from{opacity:0}to{opacity:1}}
.fade-up{animation:fadeUp .45s ease forwards;}
.fade-in{animation:fadeIn .35s ease forwards;}

section[data-testid="stSidebar"]{
  background:var(--surface)!important;
  border-right:1px solid var(--border)!important;
  width:268px!important;
}
section[data-testid="stSidebar"]>div{padding:0!important;}

.stSelectbox>div>div{
  background:var(--surface2)!important;border:1px solid var(--border)!important;
  border-radius:var(--r-sm)!important;color:var(--text2)!important;font-size:13px!important;
}
.stSelectbox>div>div:hover{border-color:var(--accent)!important;}
.stSelectbox label,.stRadio>label,.stCheckbox>label{
  color:var(--muted)!important;font-size:10px!important;font-weight:600!important;
  letter-spacing:.1em!important;text-transform:uppercase!important;
}
.stRadio [data-testid="stMarkdownContainer"] p,
.stCheckbox label p{color:var(--text2)!important;font-size:13px!important;}

.stButton>button{
  background:var(--grad)!important;color:#fff!important;border:none!important;
  border-radius:var(--r-sm)!important;font-family:'Inter',sans-serif!important;
  font-size:13px!important;font-weight:600!important;padding:11px 22px!important;
  transition:all .2s cubic-bezier(.4,0,.2,1)!important;
  box-shadow:0 4px 15px rgba(59,130,246,.3)!important;
}
.stButton>button:hover{
  transform:translateY(-2px)!important;
  box-shadow:0 8px 25px rgba(59,130,246,.45)!important;
}
.stButton>button:active{transform:translateY(0)!important;}

.stTextInput input{
  background:var(--surface2)!important;border:1px solid var(--border)!important;
  border-radius:var(--r-sm)!important;color:var(--text)!important;
  font-family:'Inter',sans-serif!important;font-size:14px!important;
  padding:10px 14px!important;transition:border-color .2s!important;
}
.stTextInput input:focus{border-color:var(--accent)!important;box-shadow:var(--glow)!important;}
.stTextInput label{
  color:var(--muted)!important;font-size:10px!important;font-weight:600!important;
  letter-spacing:.1em!important;text-transform:uppercase!important;
}

div[data-testid="stMetric"]{
  background:var(--surface)!important;border:1px solid var(--border)!important;
  border-radius:var(--r)!important;padding:20px 22px!important;
  position:relative!important;overflow:hidden!important;
  transition:all .25s cubic-bezier(.4,0,.2,1)!important;
}
div[data-testid="stMetric"]::after{
  content:'';position:absolute;top:0;left:0;width:2px;height:100%;background:var(--grad);
}
div[data-testid="stMetric"]:hover{
  border-color:var(--accent)!important;transform:translateY(-3px)!important;
  box-shadow:var(--glow)!important;
}
div[data-testid="stMetricLabel"] p{
  color:var(--muted)!important;font-size:10px!important;font-weight:600!important;
  letter-spacing:.1em!important;text-transform:uppercase!important;
}
div[data-testid="stMetricValue"]{
  color:var(--text)!important;font-size:28px!important;font-weight:800!important;
  letter-spacing:-.03em!important;
}

.stTabs [data-baseweb="tab-list"]{
  background:var(--surface)!important;border-radius:50px!important;
  padding:4px 5px!important;border:1px solid var(--border)!important;
  display:inline-flex!important;gap:2px!important;margin-bottom:24px!important;
}
.stTabs [data-baseweb="tab"]{
  background:transparent!important;color:var(--muted)!important;
  border-radius:50px!important;font-size:12px!important;font-weight:600!important;
  padding:8px 16px!important;border:none!important;transition:all .2s!important;
  white-space:nowrap!important;
}
.stTabs [data-baseweb="tab"]:hover{color:var(--text2)!important;}
.stTabs [aria-selected="true"]{
  background:var(--grad)!important;color:#fff!important;
  box-shadow:0 4px 15px rgba(59,130,246,.4)!important;
}
.stTabs [data-baseweb="tab-highlight"],
.stTabs [data-baseweb="tab-border"]{display:none!important;}

.stDataFrame{border:1px solid var(--border)!important;border-radius:var(--r)!important;overflow:hidden!important;}
.stSuccess{background:rgba(16,185,129,.08)!important;border:1px solid rgba(16,185,129,.25)!important;border-radius:var(--r-sm)!important;}
.stInfo{background:rgba(59,130,246,.06)!important;border:1px solid rgba(59,130,246,.2)!important;border-radius:var(--r-sm)!important;}
.stWarning{background:rgba(245,158,11,.06)!important;border:1px solid rgba(245,158,11,.2)!important;border-radius:var(--r-sm)!important;}
.stError{background:rgba(239,68,68,.08)!important;border:1px solid rgba(239,68,68,.2)!important;border-radius:var(--r-sm)!important;}
.stSpinner>div{border-top-color:var(--accent)!important;}

.iq-section{font-size:10px;font-weight:700;letter-spacing:.12em;text-transform:uppercase;
  color:var(--muted);margin:20px 0 14px;padding-bottom:10px;border-bottom:1px solid var(--border);}
.iq-card{background:var(--surface);border:1px solid var(--border);border-radius:var(--r);
  padding:22px;transition:all .25s ease;}
.iq-card:hover{border-color:var(--border2);box-shadow:var(--shadow);transform:translateY(-1px);}
.iq-card-glow{border-color:rgba(59,130,246,.2);
  background:linear-gradient(135deg,var(--surface),rgba(59,130,246,.04));}

.iq-badge{display:inline-flex;align-items:center;gap:4px;font-size:10px;font-weight:700;
  letter-spacing:.08em;text-transform:uppercase;padding:4px 10px;border-radius:50px;}
.iq-badge-blue{background:rgba(59,130,246,.12);border:1px solid rgba(59,130,246,.3);color:#93C5FD;}
.iq-badge-green{background:rgba(16,185,129,.1);border:1px solid rgba(16,185,129,.28);color:#6EE7B7;}
.iq-badge-red{background:rgba(239,68,68,.1);border:1px solid rgba(239,68,68,.28);color:#FCA5A5;}
.iq-badge-gold{background:linear-gradient(135deg,rgba(245,158,11,.15),rgba(239,68,68,.1));
  border:1px solid rgba(245,158,11,.3);color:#FCD34D;}
.iq-badge-purple{background:rgba(139,92,246,.12);border:1px solid rgba(139,92,246,.3);color:#C4B5FD;}
.iq-div{height:1px;background:var(--border);margin:20px 0;}

.iq-msg-user{display:flex;justify-content:flex-end;margin:10px 0;animation:fadeUp .3s ease;}
.iq-bubble-user{background:var(--grad);color:#fff;padding:12px 18px;
  border-radius:18px 18px 4px 18px;max-width:70%;font-size:14px;line-height:1.65;
  box-shadow:0 4px 20px rgba(59,130,246,.3);}
.iq-msg-ai{display:flex;align-items:flex-start;gap:10px;margin:10px 0;animation:fadeUp .3s ease;}
.iq-avatar{width:32px;height:32px;background:var(--grad);border-radius:50%;
  display:flex;align-items:center;justify-content:center;
  font-size:11px;font-weight:800;color:#fff;flex-shrink:0;}
.iq-bubble-ai{background:var(--surface2);border:1px solid var(--border2);color:var(--text2);
  padding:13px 18px;border-radius:4px 18px 18px 18px;max-width:76%;
  font-size:14px;line-height:1.75;}

.report-section{background:var(--surface);border:1px solid var(--border);
  border-radius:var(--r);padding:22px 26px;margin-bottom:16px;}
.report-section-title{font-size:13px;font-weight:700;color:var(--text);
  margin-bottom:12px;display:flex;align-items:center;gap:8px;}
.report-section-body{font-size:13px;color:var(--text2);line-height:1.8;}

.iq-locked-wrap{position:relative;}
.iq-lock-blur{filter:blur(4px);pointer-events:none;user-select:none;}
.iq-lock-overlay{position:absolute;inset:0;display:flex;flex-direction:column;
  align-items:center;justify-content:center;background:rgba(7,11,20,.75);
  backdrop-filter:blur(4px);border-radius:var(--r);border:1px solid var(--border2);z-index:5;gap:8px;}

.iq-usecase{background:var(--surface);border:1px solid var(--border);border-radius:var(--r);
  padding:18px 20px;transition:all .2s ease;}
.iq-usecase:hover{border-color:var(--accent);box-shadow:var(--glow);transform:translateY(-2px);}

.provider-status{display:flex;align-items:center;gap:8px;padding:10px 14px;
  background:var(--surface2);border:1px solid var(--border);
  border-radius:var(--r-sm);font-size:12px;color:var(--text2);}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
SAMPLE_CSV = """Date,Product,Region,Sales,Units,Returns,Marketing_Spend,Customer_Satisfaction
2024-01,Widget A,North,45200,312,14,8500,4.2
2024-01,Widget B,North,28900,198,8,4200,3.8
2024-01,Widget A,South,38700,267,21,7100,4.0
2024-01,Widget B,South,19400,133,5,3300,3.7
2024-02,Widget A,North,52100,359,11,9200,4.3
2024-02,Widget B,North,31200,214,9,4800,3.9
2024-02,Widget A,South,41800,288,18,7600,4.1
2024-02,Widget B,South,22700,156,6,3900,3.8
2024-03,Widget A,North,61400,423,16,10100,4.4
2024-03,Widget B,North,37500,258,12,5500,4.0
2024-03,Widget A,South,49200,339,23,8700,4.2
2024-03,Widget B,South,26800,184,7,4400,3.9"""

SAMPLE_REVIEWS = """review_id,product,review_text,rating
1,Widget A,"Absolutely love this product! Works perfectly and great value.",5
2,Widget B,"Decent quality but took too long to arrive. Packaging was damaged.",3
3,Widget A,"Outstanding! Exceeded my expectations. Will definitely buy again.",5
4,Widget B,"Product broke after one week. Very disappointed with the quality.",1
5,Widget A,"Good product overall, minor issues with setup but support helped.",4
6,Widget B,"Not worth the price. Cheaper alternatives available elsewhere.",2
7,Widget A,"Fast delivery, excellent packaging, product works as described.",5
8,Widget B,"Average product. Does the job but nothing special about it.",3"""

INDUSTRY_PROMPTS = {
    "General":   "Focus on revenue trends, performance gaps, and growth drivers.",
    "Ecommerce": "Focus on AOV, conversion rates, cart abandonment, and seasonal trends.",
    "SaaS":      "Focus on MRR, churn rate, CAC, LTV, and expansion revenue.",
    "Retail":    "Focus on sell-through rate, inventory turnover, basket size, and margin.",
    "Marketing": "Focus on ROAS, CPL, CTR, funnel conversion, and campaign ROI.",
}

PROVIDERS = {
    "groq":   {"name":"Groq",   "icon":"⚡", "model":"llama-3.3-70b-versatile",
               "prefix":"gsk_", "url":"https://console.groq.com/keys",
               "hint":"Starts with gsk_  ·  Free tier available"},
    "openai": {"name":"OpenAI", "icon":"🤖", "model":"gpt-4o-mini",
               "prefix":"sk-",  "url":"https://platform.openai.com/api-keys",
               "hint":"Starts with sk-   ·  Pay-per-use"},
    "gemini": {"name":"Gemini", "icon":"✨", "model":"gemini-3.1-flash-lite",
               "prefix":"AIza", "url":"https://aistudio.google.com/app/apikey",
               "hint":"Starts with AIza  ·  Free tier available"},
}

SECTION_ICONS = {
    "data overview":"📊","key findings":"🔍","trend analysis":"📈",
    "risk signals":"⚠️","strategic recommendations":"✅","executive takeaway":"💡",
    "recommendations":"✅","sentiment overview":"💬",
    "positive highlights":"🟢","top issues":"🔴",
}

USE_CASES = [
    ("📈","Sales Analysis",    "Revenue trends, top performers, growth opportunities"),
    ("💬","Review Sentiment",  "Analyze customer reviews and feedback automatically"),
    ("🎯","Customer Insights", "Segments, behavior patterns, satisfaction drivers"),
    ("💰","Revenue Trends",    "MRR growth, churn impact, expansion revenue"),
    ("📊","Marketing ROI",     "Campaign performance, ROAS, and attribution"),
    ("🔮","Forecasting",       "Predict future trends and model scenarios"),
]

# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
for k, v in [("messages",[]),("report_cache",None),("df",None),
             ("pending_input",None),("sentiment_cache",None),
             ("last_sent",""),("api_key",""),("provider",None),
             ("key_verified",False)]:
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────────────────────────────────────
# PROVIDER DETECTION
# ─────────────────────────────────────────────────────────────────────────────
def detect_provider(key: str):
    k = key.strip()
    if k.startswith("gsk_"):  return "groq"
    if k.startswith("sk-"):   return "openai"
    if k.startswith("AIza"):  return "gemini"
    return None

# ─────────────────────────────────────────────────────────────────────────────
# AI CALL — MULTI-PROVIDER (all fixes applied)
# ─────────────────────────────────────────────────────────────────────────────
def call_ai(system: str, user: str) -> str:
    key      = st.session_state.api_key.strip()
    provider = st.session_state.provider

    if not key or not provider:
        return "⚠️ No API key configured. Please add your key in the sidebar."

    try:
        if provider == "groq":
            from groq import Groq
            r = Groq(api_key=key).chat.completions.create(
                model=PROVIDERS["groq"]["model"], max_tokens=1800,
                messages=[{"role":"system","content":system},
                          {"role":"user",  "content":user}])
            return r.choices[0].message.content

        elif provider == "openai":
            from openai import OpenAI
            r = OpenAI(api_key=key).chat.completions.create(
                model=PROVIDERS["openai"]["model"], max_tokens=1800,
                messages=[{"role":"system","content":system},
                          {"role":"user",  "content":user}])
            return r.choices[0].message.content

        elif provider == "gemini":
            # ── Fixed: use google-generativeai (stable, stateless) ──
            import google.generativeai as genai
            genai.configure(api_key=key)
            model = genai.GenerativeModel(
                model_name="gemini-3.1-flash-lite",
                system_instruction=system
            )
            r = model.generate_content(user)
            return r.text

    except Exception as e:
        err = str(e)
        if any(w in err.lower() for w in ["auth","api_key","invalid","permission"]):
            st.session_state.key_verified = False
        return f"⚠️ API error: {err}"

def verify_key(key: str, provider: str):
    try:
        if provider == "groq":
            from groq import Groq
            Groq(api_key=key).chat.completions.create(
                model=PROVIDERS["groq"]["model"], max_tokens=5,
                messages=[{"role":"user","content":"hi"}])
        elif provider == "openai":
            from openai import OpenAI
            OpenAI(api_key=key).chat.completions.create(
                model=PROVIDERS["openai"]["model"], max_tokens=5,
                messages=[{"role":"user","content":"hi"}])
        elif provider == "gemini":
            import google.generativeai as genai
            genai.configure(api_key=key)
            genai.GenerativeModel("gemini-3.1-flash-lite").generate_content("hi")
        return True, "Key verified"
    except Exception as e:
        return False, str(e)

# ─────────────────────────────────────────────────────────────────────────────
# DATA HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def load_data(f):
    return pd.read_csv(f) if f.name.endswith(".csv") else pd.read_excel(f)

def clean_text(t):
    return re.sub(r'[^\x00-\x7F]+', '', t).strip()

def get_data_context(df, n=8):
    num = df.select_dtypes(include='number')
    return f"""DATASET: {df.shape[0]} rows x {df.shape[1]} cols
Columns: {list(df.columns)}
Missing: {df.isnull().sum().to_dict()}
Stats:\n{num.describe().round(2).to_string() if not num.empty else 'N/A'}
Sample:\n{df.head(n).to_csv(index=False)}"""

def generate_kpis(df):
    missing = int(df.isnull().sum().sum())
    return {"rows":df.shape[0],"cols":df.shape[1],"missing":missing,
            "complete":f"{100-missing/df.size*100:.1f}%",
            "num_cols":len(df.select_dtypes(include='number').columns)}

def detect_text_columns(df):
    return [c for c in df.columns if df[c].dtype == object
            and df[c].dropna().astype(str).str.len().mean() > 30]

def detect_score_columns(df):
    kw = ['rating','score','review','sentiment','stars','satisfaction']
    return [c for c in df.columns
            if any(k in c.lower() for k in kw)
            and df[c].dtype in ['float64','int64']]

def has_sentiment_data(df):
    return bool(detect_text_columns(df) or detect_score_columns(df))

# ─────────────────────────────────────────────────────────────────────────────
# SAFE BAR CHART HELPER (fixes ValueError: cannot insert column)
# ─────────────────────────────────────────────────────────────────────────────
def safe_bar_chart(df, x_col, y_col):
    """Group and render bar chart safely — avoids duplicate column errors."""
    try:
        chart_data = df.groupby(x_col)[y_col].sum()
        st.bar_chart(chart_data, use_container_width=True)
    except Exception as e:
        st.warning(f"Could not render chart: {e}")

def safe_line_chart(df, x_col, y_col):
    """Render line chart safely."""
    try:
        st.line_chart(df.set_index(x_col)[y_col], use_container_width=True)
    except Exception:
        try:
            st.line_chart(df[y_col], use_container_width=True)
        except Exception as e:
            st.warning(f"Could not render chart: {e}")

# ─────────────────────────────────────────────────────────────────────────────
# SENTIMENT ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
def run_sentiment_analysis(df):
    text_cols  = detect_text_columns(df)
    score_cols = detect_score_columns(df)
    res = {"type":None,"text_col":None,"score_col":None,
           "distribution":{},"themes_pos":[],"themes_neg":[],
           "recommendations":"","sample_df":None}

    if text_cols:
        col = text_cols[0]; res["type"] = "text"; res["text_col"] = col
        reviews = df[col].dropna().astype(str).tolist()[:40]
        sys_p = """Analyze reviews and return ONLY valid JSON, no explanation:
{"distribution":{"Positive":0,"Neutral":0,"Negative":0},
"themes_positive":["t1","t2","t3"],"themes_negative":["t1","t2","t3"],"summary":"one sentence"}"""
        raw = call_ai(sys_p, "Reviews:\n" + "\n".join(
            [f"{i+1}. {r[:200]}" for i,r in enumerate(reviews)]))
        try:
            data = json.loads(re.sub(r'```json|```','',raw).strip())
            res["distribution"] = data.get("distribution",{})
            res["themes_pos"]   = data.get("themes_positive",[])
            res["themes_neg"]   = data.get("themes_negative",[])
        except:
            res["distribution"] = {"Positive":0,"Neutral":0,"Negative":0}

        label_raw = call_ai(
            'Label each review Positive, Neutral, or Negative. Return ONLY a JSON array.',
            "Reviews:\n" + "\n".join([f"{i+1}. {r[:150]}" for i,r in enumerate(reviews)]))
        try:
            labels = json.loads(re.sub(r'```json|```','',label_raw).strip())
            sample = df[col].dropna().head(40).reset_index(drop=True)
            res["sample_df"] = pd.DataFrame({"Review":sample,"Sentiment":labels[:len(sample)]})
        except: pass

    elif score_cols:
        col = score_cols[0]; res["type"] = "scores"; res["score_col"] = col
        scores = df[col].dropna(); mx = scores.max()
        if mx <= 5:
            pos=(scores>=4).sum(); neu=((scores>=3)&(scores<4)).sum(); neg=(scores<3).sum()
        else:
            pos=(scores>=7).sum(); neu=((scores>=5)&(scores<7)).sum(); neg=(scores<5).sum()
        res["distribution"] = {"Positive":int(pos),"Neutral":int(neu),"Negative":int(neg)}

    total = sum(res["distribution"].values()) or 1
    pct   = {k: round(v/total*100) for k,v in res["distribution"].items()}
    res["recommendations"] = call_ai(
        "You are a CX consultant. Give 4 specific actionable recommendations as a numbered list with bold titles.",
        f"Sentiment: {pct}%\nPositive themes: {res['themes_pos']}\nNegative themes: {res['themes_neg']}")
    return res

# ─────────────────────────────────────────────────────────────────────────────
# REPORT WITH INLINE CHARTS
# ─────────────────────────────────────────────────────────────────────────────
def parse_report_sections(text):
    sections = []; cur_title = "Intro"; cur_body = []
    for line in text.split("\n"):
        ls = line.strip()
        if ls.startswith("## "):
            if cur_body: sections.append({"title":cur_title,"body":"\n".join(cur_body)})
            cur_title = ls[3:]; cur_body = []
        elif ls.startswith("### "):
            if cur_body: sections.append({"title":cur_title,"body":"\n".join(cur_body)})
            cur_title = ls[4:]; cur_body = []
        elif ls:
            cur_body.append(ls)
    if cur_body: sections.append({"title":cur_title,"body":"\n".join(cur_body)})
    return sections

def render_section_chart(df, title):
    """Auto-insert relevant chart after each report section."""
    num_cols = df.select_dtypes(include='number').columns.tolist()
    cat_cols = [c for c in df.columns if df[c].nunique() < 20 and df[c].dtype == object]
    date_col = next((c for c in df.columns if 'date' in c.lower()), None)
    tl = title.lower()
    if not num_cols: return

    if any(w in tl for w in ["trend","time","growth","monthly","quarterly"]):
        if date_col:
            st.caption(f"📈 {num_cols[0]} over {date_col}")
            safe_line_chart(df, date_col, num_cols[0])

    elif any(w in tl for w in ["finding","overview","performance","sales","revenue"]):
        if cat_cols:
            st.caption(f"📊 {num_cols[0]} by {cat_cols[0]}")
            safe_bar_chart(df, cat_cols[0], num_cols[0])

    elif any(w in tl for w in ["risk","return","loss","churn"]):
        risk = [c for c in num_cols if any(k in c.lower() for k in ['return','loss','risk','churn'])]
        col  = risk[0] if risk else num_cols[-1]
        if cat_cols:
            st.caption(f"⚠️ {col} by {cat_cols[0]}")
            safe_bar_chart(df, cat_cols[0], col)

def render_report_with_charts(df, text):
    for sec in parse_report_sections(text):
        tl   = sec["title"].lower()
        icon = next((v for k,v in SECTION_ICONS.items() if k in tl), "◈")
        st.markdown(f"""
        <div class="report-section">
          <div class="report-section-title">{icon} {sec['title']}</div>
          <div class="report-section-body">{sec['body'].replace(chr(10),'<br>')}</div>
        </div>""", unsafe_allow_html=True)
        if any(w in tl for w in ["finding","trend","risk","overview","performance","sales"]):
            render_section_chart(df, sec["title"])

# ─────────────────────────────────────────────────────────────────────────────
# EXPORT: MARKDOWN
# ─────────────────────────────────────────────────────────────────────────────
def generate_markdown_export(report_text, title, company="InsightIQ"):
    date  = datetime.datetime.now().strftime("%B %d, %Y")
    lines = report_text.strip().split('\n')
    ICONS = {"data overview":"📊","key findings":"🔍","trend analysis":"📈",
             "risk signals":"⚠️","strategic recommendations":"✅","executive takeaway":"💡"}
    md = [f"# {title}", f"\n> **{company}** · {date} · Confidential\n", "---\n"]
    secs = [l[4:].strip() for l in lines if l.strip().startswith('### ')]
    if secs:
        md.append("## Table of Contents\n")
        for i,s in enumerate(secs,1): md.append(f"{i}. {s}")
        md.append("\n---\n")
    for line in lines:
        ls = line.strip()
        if ls.startswith('## '): continue
        elif ls.startswith('### '):
            sec = ls[4:]; icon = ICONS.get(sec.lower(),'◈')
            md.append(f"\n## {icon} {sec}\n")
        else:
            md.append(ls)
    md.append(f"\n---\n*Generated by {company} · {date}*")
    return '\n'.join(md).encode('utf-8')

# ─────────────────────────────────────────────────────────────────────────────
# EXPORT: WORD
# ─────────────────────────────────────────────────────────────────────────────
def generate_word_export(report_text, title, company="InsightIQ"):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)
    sec.left_margin=Cm(2.5);    sec.right_margin=Cm(2.5)
    sec.top_margin=Cm(2.5);     sec.bottom_margin=Cm(2)

    NAVY=RGBColor(0x0B,0x11,0x20); ACCENT=RGBColor(0x25,0x63,0xEB)
    TEXT=RGBColor(0x1E,0x29,0x3B); TEXT2=RGBColor(0x47,0x55,0x69)
    WHITE=RGBColor(0xFF,0xFF,0xFF); MUTED=RGBColor(0x64,0x74,0x8B)
    SEC_RGB = {
        'data overview':            RGBColor(0x25,0x63,0xEB),
        'key findings':             RGBColor(0x7C,0x3A,0xED),
        'trend analysis':           RGBColor(0x08,0x91,0xB2),
        'risk signals':             RGBColor(0xDC,0x26,0x26),
        'strategic recommendations':RGBColor(0x05,0x96,0x69),
        'executive takeaway':       RGBColor(0xD9,0x77,0x06),
        'recommendations':          RGBColor(0x05,0x96,0x69),
    }
    ICONS2 = {'data overview':'📊 ','key findings':'🔍 ','trend analysis':'📈 ',
              'risk signals':'⚠ ','strategic recommendations':'✅ ','executive takeaway':'💡 '}

    def shading(para, fill):
        pPr = para._p.get_or_add_pPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), fill);  pPr.append(shd)

    def border_bottom(para, col_hex, size=8):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        b    = OxmlElement('w:bottom')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(size))
        b.set(qn('w:space'),'4');    b.set(qn('w:color'),col_hex)
        pBdr.append(b); pPr.append(pBdr)

    # Cover
    cover = doc.add_paragraph()
    shading(cover,'0B1120')
    cover.paragraph_format.space_before=Pt(0); cover.paragraph_format.space_after=Pt(0)
    r=cover.add_run(f"\n\n{company.upper()}\n")
    r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    r2=cover.add_run(f"{title}\n")
    r2.font.size=Pt(24); r2.font.bold=True; r2.font.color.rgb=WHITE
    r3=cover.add_run(f"\n{datetime.datetime.now().strftime('%B %d, %Y')}  ·  Confidential\n\n\n")
    r3.font.size=Pt(9); r3.font.color.rgb=RGBColor(0x94,0xA3,0xB8)
    stripe=doc.add_paragraph(); shading(stripe,'2563EB')
    stripe.paragraph_format.space_before=Pt(0); stripe.paragraph_format.space_after=Pt(0)
    stripe.add_run(" ")
    doc.add_page_break()

    lines = report_text.strip().split('\n')
    cur_sec=None; cur_body=[]

    def flush(sec_title, body_lines):
        if not sec_title: return
        key   = sec_title.lower().strip()
        color = SEC_RGB.get(key, ACCENT)
        hex_c = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        icon  = ICONS2.get(key,'')

        h = doc.add_paragraph(); shading(h, hex_c)
        h.paragraph_format.space_before=Pt(12); h.paragraph_format.space_after=Pt(2)
        h.paragraph_format.left_indent=Cm(0.3)
        hr = h.add_run(f"  {icon}{sec_title.upper()}  ")
        hr.font.size=Pt(10); hr.font.bold=True; hr.font.color.rgb=WHITE

        i = 0
        while i < len(body_lines):
            line = body_lines[i].strip()
            if not line: i+=1; continue
            nm = re.match(r'^(\d+)\.\s+(.+)', line)
            if nm:
                p=doc.add_paragraph(); shading(p,'F8FAFC')
                p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_after=Pt(2)
                nr=p.add_run(f"{nm.group(1)}.  ")
                nr.font.size=Pt(10); nr.font.bold=True; nr.font.color.rgb=color
                tr=p.add_run(re.sub(r'\*\*(.+?)\*\*',r'\1',nm.group(2)))
                tr.font.size=Pt(10); tr.font.bold=True; tr.font.color.rgb=TEXT
                i+=1
                while i < len(body_lines):
                    sub = body_lines[i].strip()
                    if sub.startswith(('- ','* ')):
                        sp=doc.add_paragraph(); shading(sp,'F8FAFC')
                        sp.paragraph_format.left_indent=Cm(1.5); sp.paragraph_format.space_after=Pt(1)
                        sr=sp.add_run(f"— {re.sub(r'[^\x00-\x7F]+','',sub[2:])}")
                        sr.font.size=Pt(9.5); sr.font.color.rgb=TEXT2
                        i+=1
                    else: break
                continue
            if line.startswith(('- ','* ')):
                p=doc.add_paragraph(); shading(p,'F8FAFC')
                p.paragraph_format.left_indent=Cm(0.8); p.paragraph_format.space_after=Pt(2)
                dr=p.add_run("●  "); dr.font.size=Pt(7); dr.font.color.rgb=color
                tr=p.add_run(re.sub(r'\*\*(.+?)\*\*',r'\1',
                                    re.sub(r'[^\x00-\x7F]+','',line[2:])))
                tr.font.size=Pt(10); tr.font.color.rgb=TEXT2
                i+=1; continue
            p=doc.add_paragraph(); shading(p,'F8FAFC')
            p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_after=Pt(3)
            tr=p.add_run(re.sub(r'\*\*(.+?)\*\*',r'\1',
                                re.sub(r'[^\x00-\x7F]+','',line)))
            tr.font.size=Pt(10); tr.font.color.rgb=TEXT2
            i+=1

        ep=doc.add_paragraph(); border_bottom(ep, hex_c, 4)
        ep.paragraph_format.space_after=Pt(8)

    for line in lines:
        ls = line.strip()
        if ls.startswith('## '): continue
        elif ls.startswith('### '):
            flush(cur_sec, cur_body); cur_sec=ls[4:].strip(); cur_body=[]
        else:
            if cur_sec is not None: cur_body.append(ls)
    flush(cur_sec, cur_body)

    fp=doc.add_paragraph(); border_bottom(fp,'2563EB',4)
    fp.paragraph_format.space_before=Pt(10)
    fr=fp.add_run(f"© {datetime.datetime.now().year} {company}  ·  Confidential  ·  "
                  f"{datetime.datetime.now().strftime('%B %d, %Y')}")
    fr.font.size=Pt(8); fr.font.color.rgb=MUTED

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ─────────────────────────────────────────────────────────────────────────────
# EXPORT: PDF
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf_export(report_text, title, company="InsightIQ"):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     HRFlowable, Table, TableStyle, KeepTogether)
    from reportlab.pdfgen import canvas as rl_canvas

    buf=io.BytesIO(); W,H=A4
    C_DARK=colors.HexColor('#0B1120');   C_ACCENT=colors.HexColor('#2563EB')
    C_ACCENT2=colors.HexColor('#4F46E5');C_LIGHT=colors.HexColor('#F8FAFC')
    C_MUTED=colors.HexColor('#64748B');  C_BORDER=colors.HexColor('#E2E8F0')
    C_TEXT=colors.HexColor('#1E293B');   C_TEXT2=colors.HexColor('#475569')

    SEC_C = {
        'data overview':            colors.HexColor('#2563EB'),
        'key findings':             colors.HexColor('#7C3AED'),
        'trend analysis':           colors.HexColor('#0891B2'),
        'risk signals':             colors.HexColor('#DC2626'),
        'strategic recommendations':colors.HexColor('#059669'),
        'executive takeaway':       colors.HexColor('#D97706'),
        'recommendations':          colors.HexColor('#059669'),
    }
    SEC_N = {'data overview':'01','key findings':'02','trend analysis':'03',
             'risk signals':'04','strategic recommendations':'05',
             'executive takeaway':'06','recommendations':'05'}

    def _clean(t):
        t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
        return re.sub(r'[^\x00-\x7F]+','',t).strip()

    class HFC(rl_canvas.Canvas):
        def __init__(self,*a,**kw): super().__init__(*a,**kw); self._saved=[]
        def showPage(self): self._saved.append(dict(self.__dict__)); self._startPage()
        def save(self):
            n=len(self._saved)
            for s in self._saved: self.__dict__.update(s); self._draw(n); super().showPage()
            super().save()
        def _draw(self,total):
            self.saveState(); pg=self._pageNumber
            if pg==1:
                self.setFillColor(C_DARK); self.rect(0,H-85*mm,W,85*mm,fill=1,stroke=0)
                self.setFillColor(C_ACCENT); self.rect(0,H-88*mm,W,4*mm,fill=1,stroke=0)
                self.setFillColor(C_ACCENT); self.rect(0,H-85*mm,4*mm,85*mm,fill=1,stroke=0)
                self.setFillColor(colors.white); self.setFont('Helvetica-Bold',10)
                self.drawString(20*mm,H-18*mm,company.upper())
                self.setFont('Helvetica-Bold',24); self.drawString(20*mm,H-42*mm,title)
                self.setFillColor(colors.HexColor('#94A3B8')); self.setFont('Helvetica',9)
                self.drawString(20*mm,H-56*mm,
                    datetime.datetime.now().strftime("%B %d, %Y")+"  ·  Confidential")
                self.setFillColor(C_ACCENT2); self.setFillAlpha(0.2)
                self.circle(W-28*mm,H-42*mm,28*mm,fill=1,stroke=0); self.setFillAlpha(1)
            else:
                self.setFillColor(C_DARK); self.rect(0,H-14*mm,W,14*mm,fill=1,stroke=0)
                self.setFillColor(C_ACCENT); self.rect(0,H-15*mm,W,1.2*mm,fill=1,stroke=0)
                self.setFillColor(colors.white); self.setFont('Helvetica-Bold',8)
                self.drawString(20*mm,H-9*mm,company)
                self.setFillColor(colors.HexColor('#94A3B8')); self.setFont('Helvetica',8)
                self.drawRightString(W-20*mm,H-9*mm,title)
            self.setFillColor(C_LIGHT); self.rect(0,0,W,12*mm,fill=1,stroke=0)
            self.setFillColor(C_BORDER); self.rect(0,12*mm,W,0.4*mm,fill=1,stroke=0)
            self.setFillColor(C_MUTED); self.setFont('Helvetica',8)
            self.drawString(20*mm,4*mm,
                f"© {datetime.datetime.now().year} {company}  ·  Confidential")
            self.drawRightString(W-20*mm,4*mm,f"Page {pg} of {total}")
            self.setFillColor(C_ACCENT); self.circle(W/2,5.5*mm,1.2*mm,fill=1,stroke=0)
            self.restoreState()

    s_num=ParagraphStyle('sn',fontName='Helvetica-Bold',fontSize=8,
                          textColor=colors.white,alignment=TA_CENTER)
    s_ht =ParagraphStyle('ht',fontName='Helvetica-Bold',fontSize=13,
                          textColor=C_TEXT,spaceBefore=2,spaceAfter=2,leading=17)
    s_b  =ParagraphStyle('b', fontName='Helvetica',fontSize=10,
                          textColor=C_TEXT2,leading=16,spaceAfter=4)
    s_n2 =ParagraphStyle('n2',fontName='Helvetica-Bold',fontSize=10,
                          textColor=C_TEXT,leading=16,leftIndent=4,spaceAfter=2)
    s_sb =ParagraphStyle('sb',fontName='Helvetica',fontSize=10,
                          textColor=C_TEXT2,leading=15,leftIndent=14,spaceAfter=3)

    def _block(sec_title, body_lines):
        key   = sec_title.lower().strip()
        color = SEC_C.get(key, C_ACCENT)
        num   = SEC_N.get(key,'◈')
        badge = Table([[Paragraph(num,s_num)]],colWidths=[6.5*mm],rowHeights=[6.5*mm])
        badge.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,-1),color),('ALIGN',(0,0),(-1,-1),'CENTER'),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),0),('BOTTOMPADDING',(0,0),(-1,-1),0)]))
        hdr=Table([[badge,Paragraph(sec_title,s_ht)]],colWidths=[9*mm,W-40*mm-9*mm])
        hdr.setStyle(TableStyle([
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('LEFTPADDING',(0,0),(0,0),0),('LEFTPADDING',(1,0),(1,0),7),
            ('TOPPADDING',(0,0),(-1,-1),0),('BOTTOMPADDING',(0,0),(-1,-1),0)]))
        div=HRFlowable(width="100%",thickness=0.8,color=color,spaceAfter=8,spaceBefore=5)
        blk=[hdr,div]; i=0
        while i<len(body_lines):
            line=body_lines[i].strip()
            if not line: i+=1; continue
            nm=re.match(r'^(\d+)\.\s+(.+)',line)
            if nm:
                blk.append(Paragraph(f"<b>{nm.group(1)}.</b>  {_clean(nm.group(2))}",s_n2))
                i+=1
                while i<len(body_lines):
                    sub=body_lines[i].strip()
                    if sub.startswith(('- ','* ')):
                        blk.append(Paragraph(
                            f"<font color='#94A3B8'>  —</font>  {_clean(sub[2:])}",s_sb))
                        i+=1
                    else: break
                continue
            if line.startswith(('- ','* ')):
                dot=Table([['']], colWidths=[2*mm],rowHeights=[2*mm])
                dot.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),color),
                    ('TOPPADDING',(0,0),(-1,-1),0),('BOTTOMPADDING',(0,0),(-1,-1),0)]))
                row=Table([[dot,Paragraph(_clean(line[2:]),s_b)]],
                          colWidths=[5*mm,W-40*mm-5*mm])
                row.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'TOP'),
                    ('LEFTPADDING',(0,0),(-1,-1),0),
                    ('TOPPADDING',(0,0),(-1,-1),1),('BOTTOMPADDING',(0,0),(-1,-1),1)]))
                blk.append(row); i+=1; continue
            blk.append(Paragraph(_clean(line),s_b)); i+=1
        blk.append(Spacer(1,3*mm))
        card=Table([[blk]],colWidths=[W-40*mm])
        card.setStyle(TableStyle([
            ('BOX',(0,0),(-1,-1),0.5,C_BORDER),
            ('LEFTPADDING',(0,0),(-1,-1),12),('RIGHTPADDING',(0,0),(-1,-1),12),
            ('TOPPADDING',(0,0),(-1,-1),14),('BOTTOMPADDING',(0,0),(-1,-1),6),
            ('ROWBACKGROUNDS',(0,0),(-1,-1),[C_LIGHT])]))
        return KeepTogether([card,Spacer(1,4*mm)])

    doc=SimpleDocTemplate(buf,pagesize=A4,rightMargin=20*mm,leftMargin=20*mm,
                          topMargin=20*mm,bottomMargin=18*mm,title=title,author=company)
    story=[Spacer(1,92*mm)]
    cur_sec=None; cur_body=[]
    for line in report_text.split('\n'):
        ls=line.strip()
        if ls.startswith('## '): continue
        elif ls.startswith('### '):
            if cur_sec: story.append(_block(cur_sec,cur_body))
            cur_sec=ls[4:].strip(); cur_body=[]
        else:
            if cur_sec is not None: cur_body.append(ls)
    if cur_sec: story.append(_block(cur_sec,cur_body))
    doc.build(story,canvasmaker=HFC)
    buf.seek(0); return buf

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="padding:28px 20px 20px">
      <div style="display:flex;align-items:center;gap:12px">
        <div style="width:38px;height:38px;
          background:linear-gradient(135deg,#3B82F6,#6366F1,#8B5CF6);
          border-radius:11px;display:flex;align-items:center;justify-content:center;
          font-size:14px;font-weight:900;color:#fff;
          box-shadow:0 4px 16px rgba(59,130,246,.4);flex-shrink:0">IQ</div>
        <div>
          <div style="font-size:17px;font-weight:800;color:#F1F5F9;letter-spacing:-.02em">InsightIQ</div>
          <div style="font-size:10px;color:#3B82F6;letter-spacing:.12em;text-transform:uppercase;margin-top:2px;font-weight:600">AI ANALYST</div>
        </div>
      </div>
    </div>
    <div style="height:1px;background:#1E2D45;margin:0 16px 20px"></div>
    """, unsafe_allow_html=True)

    # API Key input
    st.markdown('<div style="padding:0 16px">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:10px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#64748B;margin-bottom:10px">AI Configuration</div>', unsafe_allow_html=True)

    api_key_input = st.text_input(
        "API Key", value=st.session_state.api_key,
        type="password", placeholder="Paste your API key…",
        label_visibility="collapsed"
    )

    if api_key_input != st.session_state.api_key:
        st.session_state.api_key      = api_key_input
        st.session_state.provider     = detect_provider(api_key_input)
        st.session_state.key_verified = False

    # Provider status
    if st.session_state.api_key:
        p = st.session_state.provider
        if p:
            info = PROVIDERS[p]
            if st.session_state.key_verified:
                st.markdown(f"""<div class="provider-status">
                  <span style="color:#10B981;font-weight:700">●</span>
                  <span>{info['icon']} {info['name']} connected</span>
                  <span style="margin-left:auto;font-size:10px;color:#475569">{info['model']}</span>
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""<div class="provider-status">
                  <span style="color:#F59E0B;font-weight:700">●</span>
                  <span>{info['icon']} {info['name']} detected</span>
                </div>""", unsafe_allow_html=True)
                if st.button("Verify Key", use_container_width=True):
                    with st.spinner("Verifying…"):
                        ok, msg = verify_key(st.session_state.api_key, p)
                    if ok:
                        st.session_state.key_verified = True; st.rerun()
                    else:
                        st.error(f"Key failed: {msg[:80]}")
        else:
            st.markdown("""<div class="provider-status" style="border-color:rgba(239,68,68,.3)">
              <span style="color:#EF4444;font-weight:700">●</span>
              <span>Unknown key format</span>
            </div>""", unsafe_allow_html=True)
    else:
        st.markdown("""<div class="provider-status">
          <span style="color:#475569;font-weight:700">○</span>
          <span style="color:#475569">No key entered</span>
        </div>""", unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('<div style="height:1px;background:#1E2D45;margin:16px 16px"></div>', unsafe_allow_html=True)

    st.markdown('<div style="padding:0 16px">', unsafe_allow_html=True)
    industry    = st.selectbox("Industry", ["General","Ecommerce","SaaS","Retail","Marketing"])
    st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)
    report_type = st.selectbox("Report Type", [
        "Executive Summary","Sales Performance",
        "Trend Analysis","Strategic Recommendations","Risk Analysis"])
    st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)
    export_fmt  = st.radio("Export Format", ["Markdown (.md)","Word (.docx)","PDF (.pdf)"])
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('<div style="height:1px;background:#1E2D45;margin:16px 16px"></div>', unsafe_allow_html=True)

    st.markdown('<div style="padding:0 16px">', unsafe_allow_html=True)
    is_premium = st.checkbox("⚡ Pro Mode")
    if is_premium:
        st.markdown('<div style="margin-top:8px"><span class="iq-badge iq-badge-gold">✦ Pro Active</span></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("""
    <div style="position:fixed;bottom:0;left:0;width:268px;
      padding:14px 20px;border-top:1px solid #1E2D45;background:#0F1624">
      <div style="font-size:11px;color:#1E3A5F;line-height:1.8">
        Bring your own key · Zero cost to host<br>
        <span style="color:#F97316">Groq</span> ·
        <span style="color:#10A37F">OpenAI</span> ·
        <span style="color:#4285F4">Gemini</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# ONBOARDING (no key)
# ─────────────────────────────────────────────────────────────────────────────
has_key = bool(st.session_state.api_key.strip())
df      = st.session_state.df

if not has_key:
    st.markdown("""
    <div class="fade-up" style="max-width:640px;margin:48px auto 0;text-align:center;padding:0 20px">
      <div style="margin-bottom:20px">
        <span class="iq-badge iq-badge-blue">✦ Bring Your Own Key · Always Free to Use</span>
      </div>
      <h1 style="font-size:46px;font-weight:900;color:#F1F5F9;line-height:1.1;
          letter-spacing:-.04em;margin-bottom:16px">
        Your AI<br>
        <span style="background:linear-gradient(135deg,#3B82F6,#6366F1,#8B5CF6);
          -webkit-background-clip:text;-webkit-text-fill-color:transparent">Business Analyst</span>
      </h1>
      <p style="font-size:17px;color:#64748B;line-height:1.65;margin-bottom:40px">
        Connect your free AI API key to get started.<br>
        Your key stays in your browser — never stored.
      </p>
    </div>
    """, unsafe_allow_html=True)

    cols = st.columns(3)
    for i,(pid,info) in enumerate(PROVIDERS.items()):
        with cols[i]:
            st.markdown(f"""
            <div class="iq-card" style="text-align:center;padding:24px 16px">
              <div style="font-size:28px;margin-bottom:10px">{info['icon']}</div>
              <div style="font-size:15px;font-weight:700;color:#F1F5F9;margin-bottom:6px">{info['name']}</div>
              <div style="font-size:11px;color:#64748B;margin-bottom:14px;line-height:1.5">{info['hint']}</div>
              <a href="{info['url']}" target="_blank" style="
                display:inline-block;padding:8px 18px;
                background:linear-gradient(135deg,#3B82F6,#6366F1);
                color:#fff;border-radius:8px;font-size:12px;font-weight:600;
                text-decoration:none;box-shadow:0 4px 12px rgba(59,130,246,.3)">
                Get Free Key →
              </a>
            </div>""", unsafe_allow_html=True)

    st.markdown("""
    <div style="max-width:560px;margin:32px auto 0;padding:0 20px">
      <div class="iq-card">
        <div style="font-size:13px;font-weight:700;color:#F1F5F9;margin-bottom:16px">How to get started</div>
        <div style="display:flex;flex-direction:column;gap:12px">
          <div style="display:flex;align-items:flex-start;gap:12px">
            <div style="width:24px;height:24px;background:var(--grad);border-radius:50%;
              display:flex;align-items:center;justify-content:center;
              font-size:11px;font-weight:700;color:#fff;flex-shrink:0">1</div>
            <div>
              <div style="font-size:13px;font-weight:600;color:#CBD5E1">Get a free API key</div>
              <div style="font-size:12px;color:#64748B;margin-top:2px">Click any provider above — Groq is fastest and free</div>
            </div>
          </div>
          <div style="display:flex;align-items:flex-start;gap:12px">
            <div style="width:24px;height:24px;background:var(--grad);border-radius:50%;
              display:flex;align-items:center;justify-content:center;
              font-size:11px;font-weight:700;color:#fff;flex-shrink:0">2</div>
            <div>
              <div style="font-size:13px;font-weight:600;color:#CBD5E1">Paste it in the sidebar</div>
              <div style="font-size:12px;color:#64748B;margin-top:2px">App auto-detects Groq / OpenAI / Gemini</div>
            </div>
          </div>
          <div style="display:flex;align-items:flex-start;gap:12px">
            <div style="width:24px;height:24px;background:var(--grad);border-radius:50%;
              display:flex;align-items:center;justify-content:center;
              font-size:11px;font-weight:700;color:#fff;flex-shrink:0">3</div>
            <div>
              <div style="font-size:13px;font-weight:600;color:#CBD5E1">Upload data & analyse</div>
              <div style="font-size:12px;color:#64748B;margin-top:2px">CSV, Excel, or use built-in sample datasets</div>
            </div>
          </div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""<div style="max-width:860px;margin:32px auto 0;padding:0 20px">
      <div style="font-size:10px;font-weight:700;letter-spacing:.12em;text-transform:uppercase;
          color:#475569;text-align:center;margin-bottom:16px">What you can analyse</div>
    </div>""", unsafe_allow_html=True)
    uc_cols = st.columns(3)
    for i,(icon,title,desc) in enumerate(USE_CASES):
        with uc_cols[i%3]:
            st.markdown(f"""<div class="iq-usecase">
              <div style="font-size:22px;margin-bottom:8px">{icon}</div>
              <div style="font-size:13px;font-weight:600;color:var(--text);margin-bottom:4px">{title}</div>
              <div style="font-size:12px;color:var(--muted);line-height:1.5">{desc}</div>
            </div>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────
else:
    if df is None:
        st.markdown("""
        <div class="fade-up" style="max-width:640px;margin:48px auto 0;text-align:center;padding:0 20px">
          <h1 style="font-size:40px;font-weight:900;color:#F1F5F9;letter-spacing:-.04em;margin-bottom:12px">
            Upload your data</h1>
          <p style="font-size:16px;color:#64748B;margin-bottom:32px">
            CSV or Excel · Sales data · Customer reviews · Any business dataset</p>
        </div>""", unsafe_allow_html=True)

        cl,cc,cr = st.columns([1,2.2,1])
        with cc:
            uploaded_files = st.file_uploader("⬆ Drag & drop CSV or Excel",
                type=["csv","xlsx","xls"], accept_multiple_files=True,
                label_visibility="visible")
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            s1,s2 = st.columns(2)
            with s1:
                if st.button("📊 Sample Sales Data", use_container_width=True):
                    st.session_state.df=pd.read_csv(io.StringIO(SAMPLE_CSV))
                    st.session_state.messages=[]; st.session_state.report_cache=None
                    st.session_state.sentiment_cache=None; st.rerun()
            with s2:
                if st.button("💬 Sample Reviews Data", use_container_width=True):
                    st.session_state.df=pd.read_csv(io.StringIO(SAMPLE_REVIEWS))
                    st.session_state.messages=[]; st.session_state.report_cache=None
                    st.session_state.sentiment_cache=None; st.rerun()

        if uploaded_files:
            sel = uploaded_files[0]
            if len(uploaded_files) > 1:
                chosen = st.selectbox("Select file",[f.name for f in uploaded_files])
                sel = next(f for f in uploaded_files if f.name==chosen)
            try:
                with st.spinner("Loading…"):
                    st.session_state.df = load_data(sel)
                st.session_state.messages=[]; st.session_state.report_cache=None
                st.session_state.sentiment_cache=None; st.rerun()
            except Exception as e:
                st.error(f"Could not read file: {e}")

    else:
        kpis          = generate_kpis(df)
        has_sentiment = has_sentiment_data(df)
        prov          = st.session_state.provider
        prov_info     = PROVIDERS.get(prov,{}) if prov else {}

        # Top bar
        tl, tr = st.columns([3,1])
        with tl:
            pb = (f'<span class="iq-badge iq-badge-blue" style="margin-right:6px">'
                  f'{prov_info.get("icon","◈")} {prov_info.get("name","AI")}</span>'
                  if prov_info else '')
            sb = '<span class="iq-badge iq-badge-purple">💬 Sentiment Ready</span>' if has_sentiment else ''
            xb = '<span class="iq-badge iq-badge-gold" style="margin-left:6px">✦ Pro</span>' if is_premium else ''
            st.markdown(f"""
            <div class="fade-in" style="padding:20px 0 4px">
              <div style="display:flex;align-items:center;gap:8px;margin-bottom:4px">
                <div style="font-size:24px;font-weight:800;color:#F1F5F9;letter-spacing:-.03em">Dashboard</div>
                <span class="iq-badge iq-badge-green">● Live</span>{pb}{sb}{xb}
              </div>
              <div style="font-size:13px;color:#64748B">
                {df.shape[0]:,} rows · {df.shape[1]} columns ·
                <span style="color:#3B82F6;font-weight:500">{industry}</span> mode
              </div>
            </div>""", unsafe_allow_html=True)
        with tr:
            st.markdown("<div style='padding-top:20px'></div>", unsafe_allow_html=True)
            if st.button("↑ New Dataset", use_container_width=True):
                st.session_state.df=None; st.session_state.messages=[]
                st.session_state.report_cache=None; st.session_state.sentiment_cache=None
                st.rerun()

        # KPIs
        c1,c2,c3,c4,c5 = st.columns(5)
        c1.metric("Total Rows",   f"{kpis['rows']:,}")
        c2.metric("Columns",      kpis['cols'])
        c3.metric("Missing",      kpis['missing'])
        c4.metric("Completeness", kpis['complete'])
        c5.metric("Numeric Cols", kpis['num_cols'])
        st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

        # Dynamic tabs
        tab_labels = ["  📊  Overview  ","  📈  Insights  "]
        if has_sentiment: tab_labels.append("  💬  Sentiment  ")
        tab_labels += ["  🤖  AI Analyst  ","  📋  Reports  ","  ⬇️  Export  "]
        all_tabs = st.tabs(tab_labels)
        ti = 0

        # ── OVERVIEW ──────────────────────────────────────────────────────────
        with all_tabs[ti]:
            ti += 1
            st.markdown('<div class="iq-section">Data Preview</div>', unsafe_allow_html=True)
            st.dataframe(df.head(5), use_container_width=True)
            ca,cb = st.columns(2)
            with ca:
                st.markdown('<div class="iq-section" style="margin-top:20px">Column Schema</div>', unsafe_allow_html=True)
                st.dataframe(pd.DataFrame({
                    "Column":   df.columns,
                    "Type":     df.dtypes.astype(str).values,
                    "Non-Null": df.count().values,
                    "Missing":  df.isnull().sum().values,
                }), use_container_width=True, hide_index=True)
            with cb:
                st.markdown('<div class="iq-section" style="margin-top:20px">Numeric Summary</div>', unsafe_allow_html=True)
                num = df.select_dtypes(include='number')
                st.dataframe(num.describe().round(2) if not num.empty
                             else pd.DataFrame({"info":["No numeric columns"]}),
                             use_container_width=True)

        # ── INSIGHTS ──────────────────────────────────────────────────────────
        with all_tabs[ti]:
            ti += 1
            num_cols = df.select_dtypes(include='number').columns.tolist()
            cat_cols = [c for c in df.columns if df[c].nunique()<25 and df[c].dtype==object]
            date_col = next((c for c in df.columns if 'date' in c.lower()), None)

            if not num_cols:
                st.warning("No numeric columns for visualization.")
            else:
                v1,v2,v3 = st.tabs(["  Bar Chart  ","  Line Chart  ","  Correlation  "])
                with v1:
                    r1,r2 = st.columns(2)
                    y = r1.selectbox("Metric (Y)", num_cols, key="by")
                    x = r2.selectbox("Group by",  cat_cols if cat_cols else num_cols, key="bx")
                    safe_bar_chart(df, x, y)

                with v2:
                    r1,r2 = st.columns(2)
                    ly = r1.selectbox("Metric", num_cols, key="ly")
                    lx = r2.selectbox("X axis", ([date_col]+num_cols) if date_col else num_cols, key="lx")
                    safe_line_chart(df, lx, ly)

                with v3:
                    corr = df[num_cols].corr().round(2)
                    try:
                        st.dataframe(
                            corr.style.background_gradient(cmap="RdYlGn",axis=None).format("{:.2f}"),
                            use_container_width=True)
                    except Exception:
                        st.dataframe(corr, use_container_width=True)
                    st.caption("Green = positive · Red = negative correlation")

        # ── SENTIMENT ─────────────────────────────────────────────────────────
        if has_sentiment:
            with all_tabs[ti]:
                ti += 1
                st.markdown('<div class="iq-section">Sentiment Analysis</div>', unsafe_allow_html=True)
                tc = detect_text_columns(df); sc = detect_score_columns(df)
                if tc: st.markdown(f'<span class="iq-badge iq-badge-purple">💬 Text reviews: {", ".join(tc)}</span>', unsafe_allow_html=True)
                if sc: st.markdown(f'<span class="iq-badge iq-badge-blue">⭐ Score columns: {", ".join(sc)}</span>', unsafe_allow_html=True)
                st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

                if st.button("🔍 Run Sentiment Analysis", type="primary"):
                    with st.spinner("Analysing reviews…"):
                        st.session_state.sentiment_cache = run_sentiment_analysis(df)

                if st.session_state.sentiment_cache:
                    res   = st.session_state.sentiment_cache
                    dist  = res["distribution"]
                    total = sum(dist.values()) or 1
                    pos_p = round(dist.get("Positive",0)/total*100)
                    neu_p = round(dist.get("Neutral", 0)/total*100)
                    neg_p = round(dist.get("Negative",0)/total*100)

                    sa,sb2,sc2 = st.columns(3)
                    sa.metric("😊 Positive", f"{pos_p}%", f"{dist.get('Positive',0)} reviews")
                    sb2.metric("😐 Neutral",  f"{neu_p}%", f"{dist.get('Neutral', 0)} reviews")
                    sc2.metric("😞 Negative", f"{neg_p}%", f"{dist.get('Negative',0)} reviews")

                    st.markdown(f"""
                    <div style="margin:16px 0">
                      <div style="display:flex;height:12px;border-radius:50px;overflow:hidden;gap:2px">
                        <div style="width:{pos_p}%;background:#10B981;border-radius:50px 0 0 50px"></div>
                        <div style="width:{neu_p}%;background:#F59E0B"></div>
                        <div style="width:{neg_p}%;background:#EF4444;border-radius:0 50px 50px 0"></div>
                      </div>
                      <div style="display:flex;gap:20px;margin-top:8px;font-size:11px">
                        <span style="color:#10B981">● Positive {pos_p}%</span>
                        <span style="color:#F59E0B">● Neutral {neu_p}%</span>
                        <span style="color:#EF4444">● Negative {neg_p}%</span>
                      </div>
                    </div>""", unsafe_allow_html=True)

                    if dist:
                        try:
                            st.bar_chart(pd.DataFrame({"Count":dist}), use_container_width=True)
                        except Exception: pass

                    if res["themes_pos"] or res["themes_neg"]:
                        th_a,th_b = st.columns(2)
                        with th_a:
                            st.markdown('<div class="iq-section">✅ What Customers Love</div>', unsafe_allow_html=True)
                            for t in res["themes_pos"]:
                                st.markdown(f'<div style="padding:8px 12px;background:rgba(16,185,129,.08);border:1px solid rgba(16,185,129,.2);border-radius:8px;font-size:13px;color:#6EE7B7;margin-bottom:6px">✓ {t}</div>', unsafe_allow_html=True)
                        with th_b:
                            st.markdown('<div class="iq-section">⚠️ Pain Points</div>', unsafe_allow_html=True)
                            for t in res["themes_neg"]:
                                st.markdown(f'<div style="padding:8px 12px;background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.2);border-radius:8px;font-size:13px;color:#FCA5A5;margin-bottom:6px">✗ {t}</div>', unsafe_allow_html=True)

                    if res.get("sample_df") is not None:
                        st.markdown('<div class="iq-section" style="margin-top:20px">Labelled Samples</div>', unsafe_allow_html=True)
                        st.dataframe(res["sample_df"], use_container_width=True, hide_index=True)

                    if res["recommendations"]:
                        st.markdown('<div class="iq-section" style="margin-top:20px">🎯 Recommendations</div>', unsafe_allow_html=True)
                        st.markdown(
                            f'<div class="iq-card iq-card-glow" style="font-size:14px;color:var(--text2);line-height:1.8">'
                            + res["recommendations"].replace('\n','<br>') + '</div>',
                            unsafe_allow_html=True)

        # ── AI ANALYST ────────────────────────────────────────────────────────
        with all_tabs[ti]:
            ti += 1
            st.markdown('<div class="iq-section">AI Analyst · Ask Anything</div>', unsafe_allow_html=True)

            sc_cols = st.columns(4)
            prompts = ["What trends do you see?","Which segment performs best?",
                       "What should I improve?","Top 3 risks in my data?"]
            for i,p in enumerate(prompts):
                if sc_cols[i].button(p, key=f"p{i}", use_container_width=True):
                    st.session_state["pending_input"] = p
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

            ic,bc = st.columns([5,1])
            with ic:
                typed = st.text_input("msg", placeholder="Ask anything about your data…",
                    label_visibility="collapsed", key="chat_text")
            with bc:
                if st.button("Send →", use_container_width=True, key="chat_send"):
                    if typed and typed != st.session_state.get("last_sent",""):
                        st.session_state["pending_input"] = typed
                        st.session_state["last_sent"] = typed

            if st.session_state.get("pending_input"):
                inp = st.session_state.pop("pending_input")
                st.session_state.messages.append({"role":"user","content":inp})
                history = "\n".join([f"{m['role'].upper()}: {m['content']}"
                                      for m in st.session_state.messages[-8:]])
                sys_p = f"""You are InsightIQ, an elite AI business analyst.
Industry: {industry}. {INDUSTRY_PROMPTS[industry]}
Answer precisely. Use bullet points. Cite actual numbers."""
                usr_p = f"Dataset:\n{get_data_context(df)}\n\nHistory:\n{history}\n\nQuestion: {inp}"
                with st.spinner("Thinking…"):
                    reply = call_ai(sys_p, usr_p)
                st.session_state.messages.append({"role":"assistant","content":reply})

            if not st.session_state.messages:
                st.markdown("""<div style="text-align:center;padding:40px 20px">
                  <div style="font-size:38px;opacity:.3;margin-bottom:12px">🤖</div>
                  <div style="font-size:14px;font-weight:600;color:#475569">Ask your first question</div>
                </div>""", unsafe_allow_html=True)
            else:
                for m in st.session_state.messages:
                    if m["role"] == "user":
                        st.markdown(f'<div class="iq-msg-user"><div class="iq-bubble-user">{m["content"]}</div></div>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<div class="iq-msg-ai"><div class="iq-avatar">IQ</div><div class="iq-bubble-ai">{m["content"].replace(chr(10),"<br>")}</div></div>', unsafe_allow_html=True)
                st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
                if st.button("Clear chat", key="clr"):
                    st.session_state.messages = []; st.rerun()

        # ── REPORTS ───────────────────────────────────────────────────────────
        with all_tabs[ti]:
            ti += 1
            st.markdown('<div class="iq-section">Report Generator · Charts Included</div>', unsafe_allow_html=True)
            focus = st.text_input("Custom focus (optional)",
                placeholder="e.g. Focus on Q1 vs Q2 regional gap")

            b1,b2 = st.columns(2)
            with b1:
                if st.button("⚡ Generate Report + Charts", type="primary", use_container_width=True):
                    sys_p = f"""You are an elite business intelligence analyst writing a {report_type}.
Industry: {industry}. {INDUSTRY_PROMPTS[industry]}
Use these exact ### section headers:
### Data Overview
### Key Findings
### Trend Analysis
### Risk Signals
### Strategic Recommendations
### Executive Takeaway
Use real numbers. Be specific."""
                    with st.spinner("Generating report…"):
                        st.session_state.report_cache = call_ai(
                            sys_p,
                            get_data_context(df) + (f"\nFocus: {focus}" if focus else ""))

            with b2:
                if is_premium:
                    if st.button("🎯 Deep Recommendations", use_container_width=True):
                        sys_p = f"""You are a McKinsey-level consultant.
Industry: {industry}. {INDUSTRY_PROMPTS[industry]}
Deliver 5 specific recommendations. Each: **Bold Title**, Problem, Action, Expected Impact, Timeline."""
                        with st.spinner("Generating…"):
                            st.session_state.report_cache = call_ai(sys_p, get_data_context(df))
                else:
                    st.markdown("""<div class="iq-locked-wrap">
                      <div class="iq-lock-blur"><div style="background:#1A2338;border-radius:14px;padding:20px;height:80px"></div></div>
                      <div class="iq-lock-overlay">
                        <span style="font-size:20px">🔒</span>
                        <div style="font-size:14px;font-weight:700;color:var(--text)">Pro Feature</div>
                        <div style="font-size:12px;color:var(--muted)">Enable Pro in sidebar</div>
                      </div></div>""", unsafe_allow_html=True)

            if st.session_state.report_cache:
                st.markdown('<div class="iq-div"></div>', unsafe_allow_html=True)
                render_report_with_charts(df, st.session_state.report_cache)

        # ── EXPORT ────────────────────────────────────────────────────────────
        with all_tabs[ti]:
            ti += 1
            st.markdown('<div class="iq-section">Export</div>', unsafe_allow_html=True)

            if not st.session_state.report_cache:
                st.info("Generate a report in the Reports tab first.")
            else:
                report = st.session_state.report_cache
                with st.expander("Preview report"):
                    st.text(report[:600] + "…" if len(report) > 600 else report)
                st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

                if export_fmt == "Markdown (.md)":
                    st.download_button("⬇️ Download as Markdown",
                        data=generate_markdown_export(report, report_type),
                        file_name="insightiq_report.md", mime="text/markdown",
                        use_container_width=True)

                elif export_fmt == "Word (.docx)":
                    try:
                        st.download_button("⬇️ Download as Word",
                            data=generate_word_export(report, report_type),
                            file_name="insightiq_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                    except ImportError:
                        st.error("pip install python-docx")

                elif export_fmt == "PDF (.pdf)":
                    try:
                        st.download_button("⬇️ Download as PDF",
                            data=generate_pdf_export(report, report_type),
                            file_name="insightiq_report.pdf", mime="application/pdf",
                            use_container_width=True)
                    except ImportError:
                        st.error("pip install reportlab")

                st.markdown('<div class="iq-div"></div>', unsafe_allow_html=True)
                st.download_button("⬇️ Export Dataset as CSV",
                    data=df.to_csv(index=False).encode("utf-8"),
                    file_name="dataset_export.csv", mime="text/csv",
                    use_container_width=True)